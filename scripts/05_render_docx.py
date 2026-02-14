#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown -> DOCX（专利初稿）
优先使用 python-docx；若不可用则回退到内置OOXML写入器（离线可用）。
"""

import argparse
import datetime as dt
from pathlib import Path
import sys
from xml.sax.saxutils import escape
import zipfile

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
VENDOR_DIR = SKILL_ROOT / ".vendor"
if VENDOR_DIR.exists():
    sys.path.insert(0, str(VENDOR_DIR))

try:
    from docx import Document
    from docx.shared import Pt, Mm, Emu
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except Exception:
    HAS_PYTHON_DOCX = False

try:
    from PIL import Image
    HAS_PIL = True
except Exception:
    HAS_PIL = False


def _insert_figures_docx(doc, figures_dir: Path):
    if not figures_dir.exists():
        return 0
    figs = sorted(figures_dir.glob("figure*.png"))
    if not figs:
        return 0

    name_map = {
        "figure1_system_arch.png": "图1  双数据源工业波形处理系统总体架构图",
        "figure2_pipeline_flow.png": "图2  数据处理方法流程图",
        "figure3_module_block.png": "图3  模块关系框图",
    }

    sec = doc.sections[0]
    max_w = int(sec.page_width - sec.left_margin - sec.right_margin)
    # 给图题与底部留白预留空间（约16mm），避免图题被挤到下一页
    max_h = int(sec.page_height - sec.top_margin - sec.bottom_margin - Mm(16))

    # 附图整体另起页，且每图单独一页
    doc.add_page_break()
    count = 0
    for i, fp in enumerate(figs, start=1):
        if i > 1:
            doc.add_page_break()

        cap = name_map.get(fp.name, f"图{i}  附图")

        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        use_w = max_w
        use_h = None
        if HAS_PIL:
            try:
                with Image.open(fp) as im:
                    iw, ih = im.size
                if iw > 0 and ih > 0:
                    scale = min(max_w / iw, max_h / ih)
                    use_w = max(1, int(iw * scale))
                    use_h = max(1, int(ih * scale))
            except Exception:
                use_w = max_w
                use_h = None

        if use_h is None:
            pic_p.add_run().add_picture(str(fp), width=Emu(use_w))
        else:
            pic_p.add_run().add_picture(str(fp), width=Emu(use_w), height=Emu(use_h))

        cp = doc.add_paragraph(cap)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rr in cp.runs:
            rr.font.name = "宋体"
            rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            rr.font.size = Pt(12)
        count += 1
    return count


def render_with_python_docx(lines, out_path: Path, figures_dir: Path | None):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Mm(297)
    sec.page_width = Mm(210)
    sec.top_margin = Mm(25)
    sec.left_margin = Mm(25)
    sec.right_margin = Mm(15)
    sec.bottom_margin = Mm(15)

    st = doc.styles["Normal"]
    st.font.name = "宋体"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for line in lines:
        # 兼容带BOM的UTF-8文本，避免首行标题无法识别
        line = line.lstrip("\ufeff").rstrip()
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(16)
            doc.add_paragraph("")
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[3:].strip())
            r.bold = True
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(14)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(12)
            continue
        if not line.strip():
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing = 1.5
        for r in p.runs:
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(12)

    fig_count = 0
    if figures_dir:
        fig_count = _insert_figures_docx(doc, figures_dir)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return fig_count


def _w_p(text: str, size_half_pt: int = 24, bold: bool = False, center: bool = False, line_15: bool = True) -> str:
    jc = '<w:jc w:val="center"/>' if center else ""
    spacing = '<w:spacing w:line="360" w:lineRule="auto"/>' if line_15 else ""
    b = "<w:b/>" if bold else ""
    # eastAsia指定宋体，ascii/hAnsi 使用 Times New Roman 兼容英文显示
    run_pr = (
        f"<w:rPr>{b}<w:rFonts w:eastAsia=\"宋体\" w:ascii=\"Times New Roman\" "
        f"w:hAnsi=\"Times New Roman\"/><w:sz w:val=\"{size_half_pt}\"/></w:rPr>"
    )
    if text == "":
        return f"<w:p><w:pPr>{jc}{spacing}</w:pPr></w:p>"
    return f"<w:p><w:pPr>{jc}{spacing}</w:pPr><w:r>{run_pr}<w:t xml:space=\"preserve\">{escape(text)}</w:t></w:r></w:p>"


def render_with_builtin_ooxml(lines, out_path: Path, figures_dir: Path | None):
    paragraphs = []
    for raw in lines:
        # 兼容带BOM的UTF-8文本，避免首行标题无法识别
        line = raw.lstrip("\ufeff").rstrip()
        if line.startswith("# "):
            paragraphs.append(_w_p(line[2:].strip(), size_half_pt=32, bold=True, center=True, line_15=False))
            paragraphs.append(_w_p("", line_15=False))
        elif line.startswith("## "):
            paragraphs.append(_w_p(line[3:].strip(), size_half_pt=28, bold=True, center=True, line_15=False))
        elif line.startswith("### "):
            paragraphs.append(_w_p(line[4:].strip(), size_half_pt=24, bold=True, center=False, line_15=True))
        elif not line.strip():
            paragraphs.append(_w_p("", line_15=False))
        else:
            paragraphs.append(_w_p(line, size_half_pt=24, bold=False, center=False, line_15=True))

    # A4 + 页边距（twips）：25mm≈1417, 15mm≈850
    sect_pr = (
        "<w:sectPr>"
        "<w:pgSz w:w=\"11906\" w:h=\"16838\"/>"
        "<w:pgMar w:top=\"1417\" w:right=\"850\" w:bottom=\"850\" w:left=\"1417\" w:header=\"708\" w:footer=\"708\" w:gutter=\"0\"/>"
        "<w:cols w:space=\"708\"/>"
        "<w:docGrid w:linePitch=\"312\"/>"
        "</w:sectPr>"
    )
    now = dt.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>
"""
    doc_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>
"""
    app_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
 xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>Codex Patent Skill</Application>
</Properties>
"""
    core_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/"
 xmlns:dcterms="http://purl.org/dc/terms/"
 xmlns:dcmitype="http://purl.org/dc/dcmitype/"
 xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>Patent Draft</dc:title>
  <dc:creator>Codex</dc:creator>
  <cp:lastModifiedBy>Codex</cp:lastModifiedBy>
  <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>
"""

    if figures_dir and figures_dir.exists():
        fig_list = sorted(figures_dir.glob("figure*.png"))
        if fig_list:
            paragraphs.append(_w_p("", line_15=False))
            paragraphs.append(_w_p("附图（请使用 python-docx 渲染版本查看内嵌图片）", size_half_pt=24, bold=True, center=False, line_15=True))
            for i, f in enumerate(fig_list, start=1):
                paragraphs.append(_w_p(f"图{i} 文件：{f.name}", size_half_pt=24, bold=False, center=False, line_15=True))

    body_xml = "".join(paragraphs)
    document_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:wpc=\"http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas\" "
        "xmlns:mc=\"http://schemas.openxmlformats.org/markup-compatibility/2006\" "
        "xmlns:o=\"urn:schemas-microsoft-com:office:office\" "
        "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\" "
        "xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\" "
        "xmlns:v=\"urn:schemas-microsoft-com:vml\" "
        "xmlns:wp14=\"http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing\" "
        "xmlns:wp=\"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing\" "
        "xmlns:w10=\"urn:schemas-microsoft-com:office:word\" "
        "xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
        "xmlns:w14=\"http://schemas.microsoft.com/office/word/2010/wordml\" "
        "xmlns:wpg=\"http://schemas.microsoft.com/office/word/2010/wordprocessingGroup\" "
        "xmlns:wpi=\"http://schemas.microsoft.com/office/word/2010/wordprocessingInk\" "
        "xmlns:wne=\"http://schemas.microsoft.com/office/word/2006/wordml\" "
        "xmlns:wps=\"http://schemas.microsoft.com/office/word/2010/wordprocessingShape\" "
        "mc:Ignorable=\"w14 wp14\">"
        f"<w:body>{body_xml}{sect_pr}</w:body></w:document>"
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("docProps/app.xml", app_xml)
        zf.writestr("docProps/core.xml", core_xml)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="inp", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--figures-dir", default="")
    args = ap.parse_args()

    # utf-8-sig可自动去除BOM，避免markdown标题渲染失败
    lines = Path(args.inp).read_text(encoding="utf-8-sig", errors="ignore").splitlines()
    outp = Path(args.out)
    figs = Path(args.figures_dir) if args.figures_dir else None

    if HAS_PYTHON_DOCX:
        fig_count = render_with_python_docx(lines, outp, figs)
    else:
        fig_count = 0
        render_with_builtin_ooxml(lines, outp, figs)

    print(f"Wrote {outp}")
    print(f"renderer={'python-docx' if HAS_PYTHON_DOCX else 'builtin-ooxml'}")
    print(f"figures_embedded={fig_count}")


if __name__ == "__main__":
    main()
