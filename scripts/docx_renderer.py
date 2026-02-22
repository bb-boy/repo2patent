#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re

WORKFLOW_NOISE_PATTERNS = [
    re.compile(r"\[(ok|warn|error)\]", re.IGNORECASE),
    re.compile(r"\bCodex\b", re.IGNORECASE),
    re.compile(r"\bagent\b", re.IGNORECASE),
    re.compile(
        r"\b(repo_fetcher|repo_indexer|evidence_builder|query_builder|patent_search|patent_fetch_claims|novelty_matrix|disclosure_builder|run_report_builder|docx_renderer)(\.py)?\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(min_unique_patents|min_agent_queries|claims_ok_ratio|fail-on-low-recall|fail-on-empty|require-min-ok-ratio|topk)\b",
        re.IGNORECASE,
    ),
]

def ensure_docx():
    try:
        from docx import Document  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        return Document, qn
    except Exception as e:
        raise RuntimeError("python-docx not installed. Install with: pip install python-docx") from e

def set_run_font(run, font_name: str, qn) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)

def apply_paragraph_font(paragraph, font_name: str, qn) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, font_name, qn)

def apply_document_style_font(doc, font_name: str, qn) -> None:
    style_names = ["Normal", "List Bullet", "List Number"]
    style_names.extend([f"Heading {i}" for i in range(1, 10)])
    for style_name in style_names:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)

def render_from_markdown(md_text: str, output_path: str, font_name: str) -> None:
    Document, qn = ensure_docx()
    doc = Document()
    apply_document_style_font(doc, font_name, qn)

    lines = md_text.splitlines()
    in_code = False
    ordered_list_counter = 0
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            ordered_list_counter = 0
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(line)
            apply_paragraph_font(p, font_name, qn)
            i += 1
            continue
        if not stripped:
            ordered_list_counter = 0
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            ordered_list_counter = 0
            level = min(len(m.group(1)), 6)
            text = m.group(2).strip()
            p = doc.add_heading(text, level=level)
            apply_paragraph_font(p, font_name, qn)
            i += 1
            continue

        if line.startswith("> "):
            ordered_list_counter = 0
            p = doc.add_paragraph(line[2:].strip())
            apply_paragraph_font(p, font_name, qn)
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", line):
            ordered_list_counter = 0
            text = re.sub(r"^\s*[-*]\s+", "", line).strip()
            p = doc.add_paragraph(text, style="List Bullet")
            apply_paragraph_font(p, font_name, qn)
            i += 1
            continue

        # Ordered list rendering:
        # 1) support "13." alone + content on next line
        # 2) renumber per contiguous list block to avoid cross-section continuation
        num_match = re.match(r"^\s*(\d+)\.\s*(.*)$", line)
        if num_match:
            ordered_list_counter += 1
            text = num_match.group(2).strip()
            if not text and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and not re.match(r"^\s*(\d+)\.\s*(.*)$", next_line):
                    text = next_line
                    i += 1
            p = doc.add_paragraph(f"{ordered_list_counter}. {text}".strip())
            apply_paragraph_font(p, font_name, qn)
            i += 1
            continue

        ordered_list_counter = 0
        p = doc.add_paragraph(stripped)
        apply_paragraph_font(p, font_name, qn)
        i += 1

    doc.save(output_path)

def detect_workflow_noise(md_text: str) -> list[str]:
    bad: list[str] = []
    for raw_line in md_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        for pat in WORKFLOW_NOISE_PATTERNS:
            if pat.search(line):
                bad.append(line)
                break
    seen = set()
    uniq: list[str] = []
    for x in bad:
        if x in seen:
            continue
        seen.add(x)
        uniq.append(x)
    return uniq[:8]

def main() -> int:
    p = argparse.ArgumentParser(description="Render disclosure markdown to .docx")
    p.add_argument("--input", "-i", required=True, help="Input disclosure.md")
    p.add_argument("--output", "-o", required=True, help="Output .docx path")
    p.add_argument("--font-name", default="\u5b8b\u4f53", help="Output document font (default: Songti)")
    args = p.parse_args()

    with open(args.input, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    if args.input.lower().endswith(".json"):
        raise SystemExit(
            "JSON input is not supported in docx_renderer. "
            "Run scripts/disclosure_builder.py first to generate disclosure.md."
        )

    noisy_lines = detect_workflow_noise(text)
    if noisy_lines:
        preview = "\n".join([f"- {x}" for x in noisy_lines])
        raise SystemExit(
            "Detected workflow/log content in disclosure markdown. "
            "Keep workflow details in run_report.md and regenerate disclosure.md from structured disclosure.json.\n"
            f"Examples:\n{preview}"
        )

    render_from_markdown(text, args.output, args.font_name)

    print(f"[ok] written: {args.output}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())

